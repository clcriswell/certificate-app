import streamlit as st
import os
import json
import tempfile
from datetime import datetime
import re
from dateutil import parser as date_parser
from pathlib import Path
from pdfminer.high_level import extract_text
import openai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import random

client = openai.OpenAI()
OPENAI_MODEL = "gpt-4o"

# Compatibility wrapper for Streamlit rerun functionality
def safe_rerun():
    """Trigger a rerun across Streamlit versions."""
    if hasattr(st, "rerun"):
        st.rerun()
    elif hasattr(st, "experimental_rerun"):
        st.experimental_rerun()

def format_certificate_date(raw_date_str):
    try:
        dt = datetime.strptime(raw_date_str, "%B %d, %Y")
    except ValueError:
        for fmt in ("%m/%d/%Y", "%Y-%m-%d"):
            try:
                dt = datetime.strptime(raw_date_str, fmt)
                break
            except ValueError:
                continue
        else:
            return "Dated ______"
    day = dt.day
    suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    month = dt.strftime("%B")
    year_map = {
        2023: "Two Thousand and Twenty-Three",
        2024: "Two Thousand and Twenty-Four",
        2025: "Two Thousand and Twenty-Five",
        2026: "Two Thousand and Twenty-Six",
        2027: "Two Thousand and Twenty-Seven",
        2028: "Two Thousand and Twenty-Eight",
        2029: "Two Thousand and Twenty-Nine",
        2030: "Two Thousand and Thirty",
    }
    year_words = year_map.get(dt.year, dt.strftime("%Y"))
    return f"Dated the {day}{suffix} of {month}\n{year_words}"

def extract_event_date(text):
    """Attempt to parse a date from freeform text."""
    patterns = [
        r"(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)[\s\t]+\d{1,2}(?:st|nd|rd|th)?(?:,?\s*\d{2,4})?",
        r"\d{1,2}[/-]\d{1,2}(?:[/-]\d{2,4})?",
        r"\d{4}-\d{2}-\d{2}",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            date_str = match.group(0)
            try:
                dt = date_parser.parse(date_str, fuzzy=True, default=datetime(datetime.today().year, 1, 1))
                if dt.year == 1900:
                    dt = dt.replace(year=datetime.today().year)
                return dt.strftime("%B %d, %Y")
            except Exception:
                continue
    return None

def determine_name_font_size(name):
    length = len(name)
    if length <= 12: return 60
    if length <= 18: return 48
    if length <= 24: return 40
    if length <= 30: return 34
    return 28

def determine_title_font_size(title):
    length = len(title)
    if length <= 20: return 28
    if length <= 30: return 24
    if length <= 36: return 20
    return 18

def format_display_title(title: str, org: str) -> str:
    """Return a human-friendly combination of title and organization."""
    title_clean = title.strip()
    org_clean = org.strip()

    generic_titles = {"organization", "committee", "organisation"}

    if not title_clean or title_clean.lower() in generic_titles or title_clean.lower() == org_clean.lower():
        return org_clean

    if title_clean and org_clean:
        return f"{title_clean} of {org_clean}"

    return title_clean or org_clean

def enhanced_commendation(name, title, org):
    base = f"On behalf of the California State Legislature, congratulations on being recognized as {title} with {org}."
    middle = "This honor reflects your dedication and the meaningful contributions you’ve made to our community."
    close = "I wish you all the best in your future endeavors."
    return f"{base} {middle} {close}"

def log_certificates(original_data, final_data, event_text, source="pasted", global_comment=""):
    log_dir = Path("logs")
    log_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().isoformat(timespec="seconds")
    log_file = log_dir / f"cert_logs_{timestamp[:10]}.jsonl"
    with log_file.open("a", encoding="utf-8") as f:
        placeholder = {
            "name": "",
            "title": "",
            "organization": "",
            "commendation": "",
        }
        for idx, final in enumerate(final_data):
            if not final.get("approved"):
                continue
            original = original_data[idx] if idx < len(original_data) else placeholder
            entry = {
                "timestamp": timestamp,
                "source": source,
                "event_text": event_text[:1000],
                "original_name": original.get("name", ""),
                "final_name": final.get("Name", ""),
                "original_title": original.get("title", ""),
                "final_title": final.get("Title", ""),
                "original_organization": original.get("organization", ""),
                "final_organization": final.get("Organization", ""),
                "original_commendation": original.get("commendation", ""),
                "final_commendation": final.get("Certificate_Text", ""),
                "approved": True,
                "reviewer_comment": final.get("reviewer_comment", ""),
                "global_comment": global_comment,
            }
            f.write(json.dumps(entry) + "\n")

def load_example_certificates(n=3):
    log_dir = Path("logs")
    if not log_dir.exists():
        return []

    entries = []
    for log_file in sorted(log_dir.glob("cert_logs_*.jsonl"), reverse=True):
        with log_file.open("r", encoding="utf-8") as f:
            for line in f:
                try:
                    cert = json.loads(line.strip())
                    if cert.get("approved"):
                        entries.append(cert)
                except json.JSONDecodeError:
                    continue
        if len(entries) >= n:
            break

    return random.sample(entries, min(len(entries), n))

def extract_certificates(event_text, event_date):
    """Call the LLM to parse certificate information from the event text."""
    cert_rows = []
    parsed_entries = []

    SYSTEM_PROMPT = f"""
You will be given the full text of a certificate request. Your task is to extract ALL individual certificates mentioned, and for each one:

- Carefully interpret the context of the event and the nature of each person's recognition
- If more than one name or organization appears in a single entry, set \"possible_split\": true
- If you're uncertain about name, title, or org, return multiple options inside \"alternatives\"
- If an organization appears to be hosting the event, omit it from the recipient's title
- Only include "title" of "organization" when someone from that organization is receiving recognition from the host

Each certificate must include:
- name
- title
- organization (if applicable)
- date_raw (or fallback to event date)
- commendation: 2–3 sentence message starting with "On behalf of the California State Legislature..." that honors their work and ends with well wishes
- optional: possible_split (true/false)
- optional: alternatives (dictionary)

The event date is: {event_date}

Return ONLY valid JSON.
"""

    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": event_text}
        ],
        temperature=0
    )

    content = response.choices[0].message.content
    cleaned = content.strip().removeprefix("```json").removesuffix("```").strip()
    parsed_entries = json.loads(cleaned)

    for parsed in parsed_entries:
        name = parsed.get("name") or "Recipient"
        title = parsed.get("title") or ""
        org = parsed.get("organization") or ""
        commendation = parsed.get("commendation") or ""

        if title.strip().lower() == "certificate of recognition":
            title = ""

        if not commendation.strip():
            commendation = enhanced_commendation(name, title, org)

        cert_rows.append({
            "Name": name,
            "Title": title,
            "Organization": org,
            "Certificate_Text": commendation,
            "Formatted_Date": format_certificate_date(parsed.get("date_raw") or event_date),
            "Tone_Category": "📝",
            "possible_split": parsed.get("possible_split", False),
            "alternatives": parsed.get("alternatives", {}),
            "Name_Size": determine_name_font_size(name),
            "Title_Size": determine_title_font_size(format_display_title(title, org)),
            "Text_Size": 14,
            "Date_Size": 12
        })

    return parsed_entries, cert_rows

def regenerate_certificate(cert, global_comment="", reviewer_comment=""):
    """Use reviewer comments to refine an existing certificate via the LLM."""
    instructions = []
    if global_comment.strip():
        instructions.append(f"Global comment: {global_comment.strip()}")
    if reviewer_comment.strip():
        instructions.append(f"Reviewer comment: {reviewer_comment.strip()}")

    if not instructions:
        return cert

    prompt = "\n".join(instructions)
    system = (
        "You update certificate details based on reviewer comments. "
        "Return ONLY valid JSON with keys name, title, organization, date_raw, commendation."
    )

    user_msg = (
        f"Current certificate:\n"
        f"Name: {cert['Name']}\n"
        f"Title: {cert['Title']}\n"
        f"Organization: {cert['Organization']}\n"
        f"Formatted_Date: {cert['Formatted_Date']}\n"
        f"Commendation: {cert['Certificate_Text']}\n\n"
        f"{prompt}"
    )

    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[{"role": "system", "content": system}, {"role": "user", "content": user_msg}],
        temperature=0
    )

    content = response.choices[0].message.content
    cleaned = content.strip().removeprefix("```json").removesuffix("```").strip()
    updated = json.loads(cleaned)

    cert["Name"] = updated.get("name", cert["Name"])
    cert["Title"] = updated.get("title", cert["Title"])
    cert["Organization"] = updated.get("organization", cert["Organization"])
    cert["Certificate_Text"] = updated.get("commendation", cert["Certificate_Text"])
    if updated.get("date_raw"):
        cert["Formatted_Date"] = format_certificate_date(updated["date_raw"])
    return cert

def split_certificate(index):
    """Split a certificate with multiple names into separate entries."""
    cert = st.session_state.cert_rows[index]
    alt_names = cert.get("alternatives", {}).get("name")
    if not alt_names:
        parts = [n.strip() for n in cert.get("Name", "").replace("&", " and ").split(" and ")]
        alt_names = [p for p in parts if p]
    if len(alt_names) < 2:
        return
    new_certs = []
    for name in alt_names[:2]:
        new_cert = cert.copy()
        new_cert["Name"] = name
        new_cert["possible_split"] = False
        new_cert["alternatives"] = {}
        new_certs.append(new_cert)
    st.session_state.cert_rows = (
        st.session_state.cert_rows[:index] + new_certs + st.session_state.cert_rows[index+1:]
    )

st.set_page_config(layout="centered")
st.title("📁 Certificate Review Assistant")

pdf_file = st.file_uploader("Upload PDF (optional)", type=["pdf"])
text_input = st.text_area("Or paste certificate request text here", height=300)

if not pdf_file and not text_input.strip():
    st.warning("Please upload a PDF or paste request text.")
    st.stop()

if pdf_file:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        tmp.write(pdf_file.read())
        tmp_path = tmp.name
    pdf_text = extract_text(tmp_path)
    os.remove(tmp_path)
    source_type = "pdf"
else:
    pdf_text = text_input
    source_type = "pasted"

# Attempt to auto-detect the event date from the text
auto_date_raw = extract_event_date(pdf_text)
if "event_date_raw" not in st.session_state:
    st.session_state.event_date_raw = auto_date_raw or ""
if not auto_date_raw:
    st.info("Event date not detected. Please enter it below.")

event_date_input = st.text_input(
    "Event Date (e.g., May 31, 2024)", value=st.session_state.event_date_raw
)
st.session_state.event_date_raw = event_date_input
event_date_raw = st.session_state.event_date_raw or datetime.today().strftime("%B %d, %Y")
formatted_event_date = format_certificate_date(event_date_raw)
st.session_state.formatted_event_date = formatted_event_date
examples = load_example_certificates(3)
few_shot_examples = ""
for idx, ex in enumerate(examples, 1):
    few_shot_examples += f"\nExample {idx}:\nName: {ex['final_name']}\nTitle: {ex['final_title']}\nOrganization: {ex['final_organization']}\nCommendation:\n{ex['final_commendation']}\n"

if "parsed_entries" not in st.session_state:
    try:
        parsed_entries, cert_rows = extract_certificates(pdf_text, event_date_raw)
    except Exception as e:
        st.error("⚠️ GPT failed to extract entries.")
        st.text(str(e))
        st.stop()

    st.session_state.parsed_entries = parsed_entries
    st.session_state.cert_rows = cert_rows
else:
    parsed_entries = st.session_state.parsed_entries
    cert_rows = st.session_state.cert_rows

for cert in cert_rows:
    cert["Formatted_Date"] = st.session_state.formatted_event_date

st.subheader("💬 Global Comments")
global_comment = st.text_area(
    "Optional: Enter general comments, tone guidance, or feedback that applies to all certificates.",
    placeholder="e.g., 'Make all commendations sound more formal.'",
    key="global_comment"
)

if st.button("🔄 Regenerate All Certificates", key="regen_all"):
    new_rows = []
    for cert in cert_rows:
        try:
            new_rows.append(regenerate_certificate(cert, global_comment, ""))
        except Exception as e:
            st.error(str(e))
            new_rows.append(cert)
    st.session_state.cert_rows = new_rows
    cert_rows = new_rows
    st.success("Certificates updated using global comment.")

st.subheader("👁 Review, Edit, and Approve Each Certificate")
final_cert_rows = []

for i, cert in enumerate(cert_rows, 1):
    display_title = format_display_title(cert['Title'], cert['Organization'])
    with st.expander(f"📜 {cert['Name']} – {display_title}"):

        if cert.get("possible_split"):
            st.warning("⚠️ This entry may include multiple recipients.")
            decision = st.radio("Would you like to split this?", ["Keep as one", "Split into two"], key=f"split_{i}")
            if decision == "Split into two" and not st.session_state.get(f"split_done_{i}"):
                split_certificate(i-1)
                st.session_state[f"split_done_{i}"] = True
                safe_rerun()

        name = st.text_input("Name", value=cert["Name"], key=f"name_{i}")
        title = st.text_input("Title", value=cert["Title"], key=f"title_{i}")
        org = st.text_input("Organization", value=cert["Organization"], key=f"org_{i}")
        text = st.text_area("📜 Commendation", cert["Certificate_Text"], height=100, key=f"text_{i}")
        name_size = determine_name_font_size(name)
        title_size = determine_title_font_size(format_display_title(title, org))
        text_size = 14
        date_size = 12
        exclude = st.checkbox("🚫 Exclude this certificate", value=False, key=f"exclude_{i}")
        approved = not exclude
        indiv_comment = st.text_area("✏️ Reviewer Comment", "", placeholder="Optional feedback on this certificate", key=f"comment_{i}")

        cert["Name"] = name
        cert["Title"] = title
        cert["Organization"] = org
        cert["Certificate_Text"] = text
        cert["Name_Size"] = name_size
        cert["Title_Size"] = title_size
        cert["Text_Size"] = text_size
        cert["Date_Size"] = date_size
        cert["approved"] = approved
        cert["reviewer_comment"] = indiv_comment

        if st.button("🔄 Regenerate Certificate", key=f"regen_{i}"):
            if indiv_comment.strip():
                try:
                    regenerate_certificate(cert, global_comment, indiv_comment)
                    cert["Name_Size"] = determine_name_font_size(cert["Name"])
                    cert["Title_Size"] = determine_title_font_size(format_display_title(cert["Title"], cert["Organization"]))
                except Exception as e:
                    st.error(str(e))
            st.session_state.cert_rows[i-1] = cert
            safe_rerun()

        st.session_state.cert_rows[i-1] = cert
        final_cert_rows.append(cert)

        st.markdown("---")
        st.markdown("#### 📄 Certificate Preview")
        lines = []
        lines.append(
            f"<div style='text-align:center; font-size:{int(name_size)}px; font-weight:bold; margin-bottom:4px;'>{name}</div>"
        )
        display_title = format_display_title(title, org)
        if display_title.strip():
            lines.append(
                f"<div style='text-align:center; font-size:{int(title_size)}px; font-weight:bold; margin-bottom:4px;'>{display_title}</div>"
            )
        lines.append(
            f"<div style='text-align:center; font-size:{int(text_size)}px; margin-top:8px;'>{text.replace(chr(10), '<br>')}</div>"
        )
        for idx, line in enumerate(cert["Formatted_Date"].split("\n")):
            mt = 20 if idx == 0 else 0
            lines.append(
                f"<div style='text-align:center; font-size:{int(date_size)}px; margin-top:{mt}px;'>{line}</div>"
            )
        lines.append(
            f"<div style='text-align:right; font-size:12px; margin-top:0;'>_____________________________________</div>"
        )
        lines.append(
            f"<div style='text-align:right; font-size:14px; margin-top:0;'>Stan Ellis</div>"
        )
        lines.append(
            f"<div style='text-align:right; font-size:14px; margin-top:0;'>Assemblyman, 32nd District</div>"
        )
        st.markdown("<br>".join(lines), unsafe_allow_html=True)


def generate_word_certificates(entries):
    doc = Document()
    base_section = doc.sections[0]
    base_section.page_height = Inches(11)
    base_section.page_width = Inches(8.5)
    base_section.top_margin = Inches(1)
    base_section.bottom_margin = Inches(0.25)
    base_section.left_margin = Inches(.75)
    base_section.right_margin = Inches(.75)

    for i, entry in enumerate(entries):
        if i > 0:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.25)
            section.left_margin = Inches(.75)
            section.right_margin = Inches(.75)

        # Initial spacer so the name block begins 4.5" from the top
        p_spacer = doc.add_paragraph()
        p_spacer.paragraph_format.space_before = Pt(225)  # 3.5" after 1" margin
        p_spacer.add_run(" ").font.size = Pt(12)

        name_size = entry.get("Name_Size", determine_name_font_size(entry["Name"]))
        title_size = max(10, round(name_size / 2))
        safe_title = determine_title_font_size(format_display_title(entry["Title"], entry["Organization"]))
        while title_size > safe_title and name_size > 20:
            name_size -= 2
            title_size = round(name_size / 2)
        text_size = max(8, round(title_size * 0.75))

        p_name = doc.add_paragraph(entry["Name"])
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.runs[0].bold = True
        p_name.runs[0].font.name = "Times New Roman"
        p_name.runs[0].font.size = Pt(name_size)
        p_name.paragraph_format.space_after = Pt(3)

        display_title = format_display_title(entry["Title"], entry["Organization"])
        if display_title.strip():
            p_title = doc.add_paragraph(display_title)
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.runs[0].bold = True
            p_title.runs[0].font.name = "Times New Roman"
            p_title.runs[0].font.size = Pt(title_size)

        p_text = doc.add_paragraph(entry["Certificate_Text"])
        p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_text.paragraph_format.space_before = Pt(10)
        p_text.runs[0].font.name = "Times New Roman"
        p_text.runs[0].font.size = Pt(text_size)

        # Spacer to position date block starting at 8.25" from the top
        spacer_gap = doc.add_paragraph()
        spacer_gap.paragraph_format.space_before = Pt(25)  # 0.5"
        spacer_gap.add_run(" ").font.size = Pt(12)

        for idx, line in enumerate(entry["Formatted_Date"].split("\n")):
            p_date = doc.add_paragraph(line)
            p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_date.paragraph_format.space_before = Pt(0 if idx > 0 else 0)
            p_date.paragraph_format.space_after = Pt(0)
            p_date.runs[0].font.name = "Times New Roman"
            p_date.runs[0].font.size = Pt(entry.get("Date_Size", 12))

        # Spacer before signature block (1.25")
        sig_spacer = doc.add_paragraph()
        sig_spacer.paragraph_format.space_before = Pt(40)
        sig_spacer.add_run(" ").font.size = Pt(12)

        for line, size in [
            ("_____________________________________", 12),
            ("Stan Ellis", 14),
            ("Assemblyman, 32nd District", 14)
        ]:
            sig = doc.add_paragraph(line)
            sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sig.paragraph_format.space_before = Pt(0)
            sig.paragraph_format.space_after = Pt(0)
            sig.runs[0].font.name = "Times New Roman"
            sig.runs[0].font.size = Pt(size)
    return doc

if st.button("📄 Generate Word Certificates"):
    approved_entries = [c for c in final_cert_rows if c.get("approved")]
    if not approved_entries:
        st.error("No certificates were approved.")
    else:
        log_certificates(parsed_entries, approved_entries, pdf_text, source="pdf" if pdf_file else "pasted", global_comment=global_comment)
        with st.spinner("Generating Word document..."):
            doc = generate_word_certificates(approved_entries)
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            doc.save(tmp.name)
            tmp.seek(0)
            st.download_button(
                label="⬇️ Download Word Certificates",
                data=tmp.read(),
                file_name="Certificates.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
