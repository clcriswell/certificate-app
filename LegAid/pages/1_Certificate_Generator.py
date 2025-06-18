import streamlit as st
import os
import json
import tempfile
from datetime import datetime
import re
from dateutil import parser as date_parser
from pathlib import Path
from pdfminer.high_level import extract_text
import openai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
import pandas as pd
from PIL import Image
import pytesseract
import random

client = openai.OpenAI()
OPENAI_MODEL = "gpt-4o"

# Font and text constraints
NAME_MIN_SIZE = 24
NAME_MAX_SIZE = 60
NAME_MAX_LINES = 1
NAME_MAX_CHARS = 35

TITLE_MIN_SIZE = 22
TITLE_MAX_SIZE = 28
TITLE_MAX_LINES = 1
TITLE_MAX_CHARS = 40

TEXT_MIN_SIZE = 20
TEXT_MAX_SIZE = 20
TEXT_MAX_LINES = 5
TEXT_MAX_CHARS = 335

# Compatibility wrapper for Streamlit rerun functionality
def safe_rerun():
    """Trigger a rerun across Streamlit versions."""
    if hasattr(st, "rerun"):
        st.rerun()
    elif hasattr(st, "experimental_rerun"):
        st.experimental_rerun()

def format_certificate_date(raw_date_str):
    try:
        dt = datetime.strptime(raw_date_str, "%B %d, %Y")
    except ValueError:
        for fmt in ("%m/%d/%Y", "%Y-%m-%d"):
            try:
                dt = datetime.strptime(raw_date_str, fmt)
                break
            except ValueError:
                continue
        else:
            return "Dated ______"
    day = dt.day
    suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    month = dt.strftime("%B")
    year_map = {
        2023: "Two Thousand and Twenty-Three",
        2024: "Two Thousand and Twenty-Four",
        2025: "Two Thousand and Twenty-Five",
        2026: "Two Thousand and Twenty-Six",
        2027: "Two Thousand and Twenty-Seven",
        2028: "Two Thousand and Twenty-Eight",
        2029: "Two Thousand and Twenty-Nine",
        2030: "Two Thousand and Thirty",
    }
    year_words = year_map.get(dt.year, dt.strftime("%Y"))
    return f"Dated the {day}{suffix} of {month}\n{year_words}"

def extract_event_date(text):
    """Attempt to parse a date from freeform text."""
    patterns = [
        r"(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)[\s\t]+\d{1,2}(?:st|nd|rd|th)?(?:,?\s*\d{2,4})?",
        r"\d{1,2}[/-]\d{1,2}(?:[/-]\d{2,4})?",
        r"\d{4}-\d{2}-\d{2}",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            date_str = match.group(0)
            try:
                dt = date_parser.parse(date_str, fuzzy=True, default=datetime(datetime.today().year, 1, 1))
                if dt.year == 1900:
                    dt = dt.replace(year=datetime.today().year)
                return dt.strftime("%B %d, %Y")
            except Exception:
                continue
    return None

def determine_name_font_size(name: str) -> int:
    """Return the optimal font size for the name."""
    return NAME_MAX_SIZE

def determine_title_font_size(title: str) -> int:
    """Return the optimal font size for the title."""
    return TITLE_MAX_SIZE if title.strip() else 0

def format_display_title(title: str, org: str) -> str:
    """Return either the title or organization depending on context."""
    title_clean = title.strip()
    org_clean = org.strip()

    generic_titles = {"organization", "committee", "organisation"}

    if not title_clean or title_clean.lower() in generic_titles or title_clean.lower() == org_clean.lower():
        return org_clean

    return title_clean or org_clean

def enhanced_commendation(name: str, title: str, org: str) -> str:
    """Return a default commendation around the TEXT_MAX_CHARS length."""
    parts = [
        "On behalf of the California State Legislature, it is my honor to recognize"
    ]
    if title and org:
        parts.append(f"your exemplary service as {title} with {org}.")
    elif title:
        parts.append(f"your exemplary service as {title}.")
    elif org:
        parts.append(f"your exemplary service with {org}.")
    else:
        parts.append("your exemplary service.")

    middle = (
        "Your tireless dedication and contributions have made a difference in our community. "
        "Your steadfast commitment to excellence sets a standard for others. "
        "May your success continue to grow and inspire others."
    )
    text = " ".join(parts) + " " + middle
    if len(text) > TEXT_MAX_CHARS:
        short_middle = (
            "Your dedication has positively impacted our community. "
            "We wish you continued success."
        )
        text = " ".join(parts) + " " + short_middle
        if len(text) > TEXT_MAX_CHARS:
            text = text[:TEXT_MAX_CHARS]
    return text

def read_uploaded_file(uploaded_file):
    """Return extracted text and source type from an uploaded file."""
    suffix = Path(uploaded_file.name).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    text = ""
    try:
        if suffix == ".pdf":
            text = extract_text(tmp_path)
        elif suffix in {".txt", ".csv"}:
            with open(tmp_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        elif suffix == ".docx":
            doc = Document(tmp_path)
            text = "\n".join(p.text for p in doc.paragraphs)
        elif suffix in {".xlsx", ".xls"}:
            df = pd.read_excel(tmp_path, header=None)
            lines = []
            for row in df.astype(str).values:
                line = " ".join(cell for cell in row if cell and cell != "nan")
                if line:
                    lines.append(line)
            text = "\n".join(lines)
        elif suffix in {".png", ".jpg", ".jpeg"}:
            try:
                text = pytesseract.image_to_string(Image.open(tmp_path))
            except Exception:
                text = ""
        return text, suffix.lstrip(".")
    finally:
        os.remove(tmp_path)

def log_certificates(original_data, final_data, event_text, source="pasted", global_comment=""):
    log_dir = Path("logs")
    log_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().isoformat(timespec="seconds")
    log_file = log_dir / f"cert_logs_{timestamp[:10]}.jsonl"
    with log_file.open("a", encoding="utf-8") as f:
        placeholder = {
            "name": "",
            "title": "",
            "organization": "",
            "commendation": "",
        }
        for idx, final in enumerate(final_data):
            if not final.get("approved"):
                continue
            original = original_data[idx] if idx < len(original_data) else placeholder
            entry = {
                "timestamp": timestamp,
                "source": source,
                "event_text": event_text[:1000],
                "original_name": original.get("name", ""),
                "final_name": final.get("Name", ""),
                "original_title": original.get("title", ""),
                "final_title": final.get("Title", ""),
                "original_organization": original.get("organization", ""),
                "final_organization": final.get("Organization", ""),
                "original_commendation": original.get("commendation", ""),
                "final_commendation": final.get("Certificate_Text", ""),
                "approved": True,
                "reviewer_comment": final.get("reviewer_comment", ""),
                "global_comment": global_comment,
            }
            f.write(json.dumps(entry) + "\n")

def load_example_certificates(n=3):
    log_dir = Path("logs")
    if not log_dir.exists():
        return []

    entries = []
    for log_file in sorted(log_dir.glob("cert_logs_*.jsonl"), reverse=True):
        with log_file.open("r", encoding="utf-8") as f:
            for line in f:
                try:
                    cert = json.loads(line.strip())
                    if cert.get("approved"):
                        entries.append(cert)
                except json.JSONDecodeError:
                    continue
        if len(entries) >= n:
            break

    return random.sample(entries, min(len(entries), n))

def extract_certificates(event_text, event_date, uniform=False):
    """Call the LLM to parse certificate information from the event text."""
    cert_rows = []
    parsed_entries = []
    template_text = ""

    if uniform:
        SYSTEM_PROMPT = f"""
You will be given the full text of a certificate request. Your task is to extract ALL individual certificates mentioned.

Return JSON with two keys:
  template: a commendation using placeholders {{name}}, {{title}}, and {{organization}}
  certificates: list of certificates each with name, title, organization (if applicable), date_raw, category, optional possible_split and alternatives

The event date is: {event_date}

Name values must be no longer than {NAME_MAX_CHARS} characters including spaces. Title values must be no longer than {TITLE_MAX_CHARS} characters including spaces. Certificate text should be around {TEXT_MAX_CHARS} characters or fewer and at most {TEXT_MAX_LINES} lines.

Return ONLY valid JSON.
"""
    else:
        SYSTEM_PROMPT = f"""
You will be given the full text of a certificate request. Your task is to extract ALL individual certificates mentioned, and for each one:

- Carefully interpret the context of the event and the nature of each person's recognition
- If more than one name or organization appears in a single entry, set \"possible_split\": true
- If you're uncertain about name, title, or org, return multiple options inside \"alternatives\"
- If an organization appears to be hosting the event, omit it from the recipient's title
- Only include "title" of "organization" when someone from that organization is receiving recognition from the host

Each certificate must include:
- name
- title
- organization (if applicable)
- date_raw (or fallback to event date)
- category: short (2–3 word) description of the recognition type
- commendation: 2–3 sentence message starting with "On behalf of the California State Legislature..." that honors their work and ends with well wishes
- optional: possible_split (true/false)
- optional: alternatives (dictionary)

The event date is: {event_date}

Name values must be no longer than {NAME_MAX_CHARS} characters including spaces. Title values must be no longer than {TITLE_MAX_CHARS} characters including spaces. Certificate text should be around {TEXT_MAX_CHARS} characters or fewer and at most {TEXT_MAX_LINES} lines.

Return ONLY valid JSON.
"""

    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": event_text}
        ],
        temperature=0
    )

    content = response.choices[0].message.content
    cleaned = content.strip().removeprefix("```json").removesuffix("```").strip()
    data = json.loads(cleaned)

    if uniform:
        template_text = data.get("template", "")
        parsed_entries = data.get("certificates", [])
    else:
        parsed_entries = data

    for parsed in parsed_entries:
        name = parsed.get("name") or "Recipient"
        title = parsed.get("title") or ""
        org = parsed.get("organization") or ""
        category = parsed.get("category", "General")
        if uniform:
            commendation = template_text
            commendation = commendation.replace("{name}", name)
            commendation = commendation.replace("{title}", title)
            commendation = commendation.replace("{organization}", org)
        else:
            commendation = parsed.get("commendation") or ""

        if title.strip().lower() == "certificate of recognition":
            title = ""

        if not commendation.strip():
            commendation = enhanced_commendation(name, title, org)

        cert_rows.append({
            "Name": name,
            "Title": title,
            "Organization": org,
            "Certificate_Text": commendation,
            "Formatted_Date": format_certificate_date(parsed.get("date_raw") or event_date),
            "Category": category,
            "Tone_Category": "📝",
            "possible_split": parsed.get("possible_split", False),
            "alternatives": parsed.get("alternatives", {}),
            "Name_Size": NAME_MAX_SIZE,
            "Title_Size": TITLE_MAX_SIZE if format_display_title(title, org).strip() else 0,
            "Text_Size": TEXT_MAX_SIZE,
            "Date_Size": 12
        })

    return parsed_entries, cert_rows, template_text

def regenerate_certificate(cert, global_comment="", reviewer_comment=""):
    """Use reviewer comments to refine an existing certificate via the LLM."""
    instructions = []
    if global_comment.strip():
        instructions.append(f"Edit All comment: {global_comment.strip()}")
    if reviewer_comment.strip():
        instructions.append(f"Reviewer comment: {reviewer_comment.strip()}")

    if not instructions:
        return cert

    prompt = "\n".join(instructions)
    system = (
        "You update certificate details based on reviewer comments. "
        f"Name must be \u2264 {NAME_MAX_CHARS} characters. Title must be \u2264 {TITLE_MAX_CHARS} characters. "
        f"Certificate text must not exceed {TEXT_MAX_CHARS} characters and {TEXT_MAX_LINES} lines. "
        "Return ONLY valid JSON with keys name, title, organization, date_raw, commendation."
    )

    user_msg = (
        f"Current certificate:\n"
        f"Name: {cert['Name']}\n"
        f"Title: {cert['Title']}\n"
        f"Organization: {cert['Organization']}\n"
        f"Formatted_Date: {cert['Formatted_Date']}\n"
        f"Commendation: {cert['Certificate_Text']}\n\n"
        f"{prompt}"
    )

    response = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[{"role": "system", "content": system}, {"role": "user", "content": user_msg}],
        temperature=0
    )

    content = response.choices[0].message.content
    cleaned = content.strip().removeprefix("```json").removesuffix("```").strip()
    updated = json.loads(cleaned)

    cert["Name"] = updated.get("name", cert["Name"])
    cert["Title"] = updated.get("title", cert["Title"])
    cert["Organization"] = updated.get("organization", cert["Organization"])
    cert["Certificate_Text"] = updated.get("commendation", cert["Certificate_Text"])
    if updated.get("date_raw"):
        cert["Formatted_Date"] = format_certificate_date(updated["date_raw"])
    return cert

def apply_global_comment(cert_rows, global_comment):
    """Apply simple global instructions to all certificates."""
    if not global_comment.strip():
        return cert_rows

    comment = global_comment.lower()

    # Set a single organization for all certificates
    org_match = re.search(r"organization(?: name)?(?: for all certificates)?\s*(?:is|=|:)\s*['\"]?([^'\"\n]+)['\"]?", comment)
    if org_match:
        org_value = org_match.group(1).strip()
        for cert in cert_rows:
            cert["Organization"] = org_value

    # Replace the entire title with the organization text
    if ("use organization instead of title" in comment or
            "replace title with organization" in comment):
        for cert in cert_rows:
            cert["Title"] = cert.get("Organization", "")
            cert["Title_Size"] = determine_title_font_size(
                format_display_title(cert["Title"], cert["Organization"])
            )

    # Replace a specific word in the title with the organization text
    replace_word = re.search(
        r"replace ['\"]?([^'\"]+)['\"]? in title with organization",
        comment,
    )
    if replace_word:
        target = replace_word.group(1).strip()
        for cert in cert_rows:
            cert["Title"] = cert.get("Title", "").replace(target, cert.get("Organization", ""))
            cert["Title_Size"] = determine_title_font_size(
                format_display_title(cert["Title"], cert["Organization"])
            )

    return cert_rows

def split_certificate(index):
    """Split a certificate with multiple names into separate entries."""
    cert = st.session_state.cert_rows[index]
    alt_names = cert.get("alternatives", {}).get("name")
    if not alt_names:
        name_field = cert.get("Name", "")
        parts = [
            n.strip()
            for n in re.split(r"\s*(?:and|&)\s*", name_field, flags=re.IGNORECASE)
        ]
        alt_names = [p for p in parts if p]
    if len(alt_names) < 2:
        return
    new_certs = []
    for name in alt_names[:2]:
        new_cert = cert.copy()
        new_cert["Name"] = name
        new_cert["possible_split"] = False
        new_cert["alternatives"] = {}
        new_certs.append(new_cert)
    st.session_state.cert_rows = (
        st.session_state.cert_rows[:index] + new_certs + st.session_state.cert_rows[index+1:]
    )
    st.session_state.expand_after_split = [index, index + 1]

st.set_page_config(layout="centered")
st.title("📁 CertCreate")

if "started" not in st.session_state:
    st.session_state.started = False

if not st.session_state.started:
    uploaded_file = st.file_uploader(
        "Upload file (optional)",
        type=["pdf", "docx", "txt", "csv", "xlsx", "xls", "png", "jpg", "jpeg"],
    )
    text_input = st.text_area(
        "Or paste certificate request text here", height=300
    )
    use_uniform = st.checkbox(
        "Keep Certificate Text Uniformed",
        value=st.session_state.get("use_uniform", False),
        key="uniform_start",
    )

    if st.button("▶️ Begin"):
        if not uploaded_file and not text_input.strip():
            st.warning("Please upload a file or paste request text before beginning.")
        else:
            if uploaded_file:
                pdf_text, source_type = read_uploaded_file(uploaded_file)
            else:
                pdf_text = text_input
                source_type = "pasted"

            st.session_state.pdf_text = pdf_text
            st.session_state.source_type = source_type
            st.session_state.use_uniform = use_uniform

            auto_date_raw = extract_event_date(pdf_text)
            st.session_state.event_date_raw = auto_date_raw or ""
            st.session_state.started = True
            safe_rerun()
    st.stop()

uploaded_file = None
text_input = None
pdf_text = st.session_state.get("pdf_text", "")
source_type = st.session_state.get("source_type", "pasted")

if st.session_state.started:
    if st.button("🔄 Start New Request"):
        for key in [
            "pdf_text",
            "source_type",
            "parsed_entries",
            "cert_rows",
            "uniform_template",
            "event_date_raw",
            "formatted_event_date",
            "use_uniform",
        ]:
            if key in st.session_state:
                del st.session_state[key]
        st.session_state.started = False
        safe_rerun()

# Attempt to auto-detect the event date from the text
auto_date_raw = extract_event_date(pdf_text)
if "event_date_raw" not in st.session_state:
    st.session_state.event_date_raw = auto_date_raw or ""
if not auto_date_raw:
    st.info("Event date not detected. Please enter it below.")

event_date_input = st.text_input(
    "Event Date (e.g., May 31, 2024)", value=st.session_state.event_date_raw
)
st.session_state.event_date_raw = event_date_input
event_date_raw = st.session_state.event_date_raw or datetime.today().strftime("%B %d, %Y")
formatted_event_date = format_certificate_date(event_date_raw)
st.session_state.formatted_event_date = formatted_event_date

use_uniform = st.session_state.get("use_uniform", False)

examples = load_example_certificates(3)
few_shot_examples = ""
for idx, ex in enumerate(examples, 1):
    few_shot_examples += (
        f"\nExample {idx}:\nName: {ex['final_name']}\nTitle: {ex['final_title']}\n"
        f"Organization: {ex['final_organization']}\nCommendation:\n{ex['final_commendation']}\n"
    )

if "parsed_entries" not in st.session_state:
    try:
        parsed_entries, cert_rows, uniform_template = extract_certificates(
            pdf_text,
            event_date_raw,
            uniform=use_uniform,
        )
    except Exception as e:
        st.error("⚠️ GPT failed to extract entries.")
        st.text(str(e))
        st.stop()

    st.session_state.parsed_entries = parsed_entries
    st.session_state.cert_rows = cert_rows
    st.session_state.uniform_template = uniform_template
else:
    parsed_entries = st.session_state.parsed_entries
    cert_rows = st.session_state.cert_rows
    uniform_template = st.session_state.get("uniform_template", "")

expanded_indices = st.session_state.pop("expand_after_split", [])

for cert in cert_rows:
    if st.session_state.event_date_raw.strip():
        cert["Formatted_Date"] = st.session_state.formatted_event_date


# Display the uniform certificate text template when enabled
if use_uniform and uniform_template:
    st.subheader("🏷️ Uniform Certificate Text")
    st.info(uniform_template)

st.subheader("💬 Edit All")
global_comment = st.text_area(
    "Optional: Enter general comments, tone guidance, or feedback that applies to all certificates.",
    placeholder="e.g., 'Make all certificates sound more formal.'",
    key="global_comment"
)

if st.button("🔄 Regenerate All Certificates", key="regen_all"):
    cert_rows = apply_global_comment(cert_rows, global_comment)
    new_rows = []
    for cert in cert_rows:
        try:
            new_rows.append(regenerate_certificate(cert, global_comment, ""))
        except Exception as e:
            st.error(str(e))
            new_rows.append(cert)
    st.session_state.cert_rows = new_rows
    cert_rows = new_rows
    st.success("Certificates updated using Edit All comment.")

st.subheader("👁 Review and Modify Individual Certificates")
final_cert_rows = []

for i, cert in enumerate(cert_rows, 1):
    display_title = format_display_title(cert['Title'], cert['Organization'])
    kwargs = {"expanded": True} if i-1 in expanded_indices else {}
    with st.expander(f"📜 {cert['Name']} – {display_title}", **kwargs):

        if cert.get("possible_split"):
            st.warning("⚠️ This entry may include multiple recipients.")
            decision = st.radio("Would you like to split this?", ["Keep as one", "Split into two"], key=f"split_{i}")
            if decision == "Split into two" and not st.session_state.get(f"split_done_{i}"):
                split_certificate(i-1)
                st.session_state[f"split_done_{i}"] = True
                safe_rerun()

        name = st.text_input(
            "Name",
            value=cert["Name"],
            key=f"name_{i}",
            max_chars=NAME_MAX_CHARS,
        )
        title = st.text_input(
            "Title",
            value=cert["Title"],
            key=f"title_{i}",
            max_chars=TITLE_MAX_CHARS,
        )
        org = st.text_input("Organization", value=cert["Organization"], key=f"org_{i}")
        text = st.text_area(
            "📜 Certificate Text",
            cert["Certificate_Text"],
            height=100,
            key=f"text_{i}",
            max_chars=TEXT_MAX_CHARS,
        )
        lines = text.splitlines()[:TEXT_MAX_LINES]
        text = "\n".join(lines)
        name_size = NAME_MAX_SIZE
        display_title = format_display_title(title, org)
        title_size = TITLE_MAX_SIZE if display_title.strip() else 0
        text_size = TEXT_MAX_SIZE
        date_size = 12
        exclude = st.checkbox("🚫 Exclude this certificate", value=False, key=f"exclude_{i}")
        approved = not exclude
        indiv_comment = st.text_area("✏️ Reviewer Comment", "", placeholder="Optional feedback on this certificate", key=f"comment_{i}")

        cert["Name"] = name
        cert["Title"] = title
        cert["Organization"] = org
        cert["Certificate_Text"] = text
        cert["Name_Size"] = NAME_MAX_SIZE
        cert["Title_Size"] = TITLE_MAX_SIZE if display_title.strip() else 0
        cert["Text_Size"] = TEXT_MAX_SIZE
        cert["Date_Size"] = date_size
        cert["approved"] = approved
        cert["reviewer_comment"] = indiv_comment

        if st.button("🔄 Regenerate Certificate", key=f"regen_{i}"):
            if indiv_comment.strip():
                try:
                    apply_global_comment([cert], global_comment)
                    regenerate_certificate(cert, global_comment, indiv_comment)
                    cert["Name_Size"] = NAME_MAX_SIZE
                    cert["Title_Size"] = TITLE_MAX_SIZE if format_display_title(cert["Title"], cert["Organization"]).strip() else 0
                except Exception as e:
                    st.error(str(e))
            st.session_state.cert_rows[i-1] = cert
            safe_rerun()

        st.session_state.cert_rows[i-1] = cert
        final_cert_rows.append(cert)

        st.markdown("---")
        st.markdown("#### 📄 Certificate Preview")
        lines = []
        lines.append(
            f"<div style='text-align:center; font-size:{int(name_size)}px; font-weight:bold; margin-bottom:4px;'>{name}</div>"
        )
        display_title = format_display_title(title, org)
        if display_title.strip():
            lines.append(
                f"<div style='text-align:center; font-size:{int(title_size)}px; font-weight:bold; margin-bottom:4px;'>{display_title}</div>"
            )
        lines.append(
            f"<div style='text-align:center; font-size:{int(text_size)}px; margin-top:8px;'>{text.replace(chr(10), '<br>')}</div>"
        )
        for idx, line in enumerate(cert["Formatted_Date"].split("\n")):
            mt = 20 if idx == 0 else 0
            lines.append(
                f"<div style='text-align:center; font-size:{int(date_size)}px; margin-top:{mt}px;'>{line}</div>"
            )
        lines.append(
            f"<div style='text-align:right; font-size:12px; margin-top:0;'>_____________________________________</div>"
        )
        lines.append(
            f"<div style='text-align:right; font-size:14px; margin-top:0;'>Stan Ellis</div>"
        )
        lines.append(
            f"<div style='text-align:right; font-size:14px; margin-top:0;'>Assemblyman, 32nd District</div>"
        )
        st.markdown("<br>".join(lines), unsafe_allow_html=True)


def generate_word_certificates(entries):
    doc = Document()
    base_section = doc.sections[0]
    base_section.page_height = Inches(11)
    base_section.page_width = Inches(8.5)
    base_section.top_margin = Inches(1)
    base_section.bottom_margin = Inches(0.25)
    base_section.left_margin = Inches(.75)
    base_section.right_margin = Inches(.75)

    for i, entry in enumerate(entries):
        if i > 0:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.25)
            section.left_margin = Inches(.75)
            section.right_margin = Inches(.75)

        # Initial spacer so the name block begins 4.5" from the top
        p_spacer = doc.add_paragraph()
        p_spacer.paragraph_format.space_before = Pt(225)  # 3.5" after 1" margin
        p_spacer.add_run(" ").font.size = Pt(12)

        name_size = NAME_MAX_SIZE
        display_title = format_display_title(entry["Title"], entry["Organization"])
        title_size = TITLE_MAX_SIZE if display_title.strip() else 0
        text_size = TEXT_MAX_SIZE

        p_name = doc.add_paragraph(entry["Name"])
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.runs[0].bold = True
        p_name.runs[0].font.name = "Times New Roman"
        p_name.runs[0].font.size = Pt(name_size)
        p_name.paragraph_format.space_after = Pt(3)

        display_title = format_display_title(entry["Title"], entry["Organization"])
        if display_title.strip():
            p_title = doc.add_paragraph(display_title)
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.runs[0].bold = True
            p_title.runs[0].font.name = "Times New Roman"
            p_title.runs[0].font.size = Pt(title_size)

        p_text = doc.add_paragraph(entry["Certificate_Text"])
        p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_text.paragraph_format.space_before = Pt(10)
        p_text.runs[0].font.name = "Times New Roman"
        p_text.runs[0].font.size = Pt(text_size)

        # Spacer to position date block starting at 8.25" from the top
        spacer_gap = doc.add_paragraph()
        spacer_gap.paragraph_format.space_before = Pt(25)  # 0.5"
        spacer_gap.add_run(" ").font.size = Pt(12)

        for idx, line in enumerate(entry["Formatted_Date"].split("\n")):
            p_date = doc.add_paragraph(line)
            p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_date.paragraph_format.space_before = Pt(0 if idx > 0 else 0)
            p_date.paragraph_format.space_after = Pt(0)
            p_date.runs[0].font.name = "Times New Roman"
            p_date.runs[0].font.size = Pt(entry.get("Date_Size", 12))

        # Spacer before signature block (1.25")
        sig_spacer = doc.add_paragraph()
        sig_spacer.paragraph_format.space_before = Pt(40)
        sig_spacer.add_run(" ").font.size = Pt(12)

    for line, size in [
            ("_____________________________________", 12),
            ("Stan Ellis", 14),
            ("Assemblyman, 32nd District", 14)
        ]:
            sig = doc.add_paragraph(line)
            sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sig.paragraph_format.space_before = Pt(0)
            sig.paragraph_format.space_after = Pt(0)
            sig.runs[0].font.name = "Times New Roman"
            sig.runs[0].font.size = Pt(size)
    return doc


def generate_pdf_certificates(entries):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    page_width, page_height = letter
    left_margin = right_margin = 0.75 * inch

    def wrap_text(text, font_name, font_size, max_width):
        lines = []
        for raw in text.split("\n"):
            words = raw.split()
            current = ""
            for word in words:
                test = f"{current} {word}".strip()
                if c.stringWidth(test, font_name, font_size) <= max_width:
                    current = test
                else:
                    if current:
                        lines.append(current)
                    current = word
            if current:
                lines.append(current)
        return lines

    for i, entry in enumerate(entries):
        if i > 0:
            c.showPage()

        name_size = NAME_MAX_SIZE
        display_title = format_display_title(entry["Title"], entry["Organization"])
        title_size = TITLE_MAX_SIZE if display_title.strip() else 0
        text_size = TEXT_MAX_SIZE
        date_size = 12

        center_x = page_width / 2
        avail_width = page_width - left_margin - right_margin

        c.setFont("Times-Bold", name_size)
        c.drawCentredString(center_x, page_height - 4.5 * inch, entry["Name"])

        if display_title.strip():
            c.setFont("Times-Bold", title_size)
            y = page_height - (4.5 + 0.25) * inch
            for line in wrap_text(display_title, "Times-Bold", title_size, avail_width):
                c.drawCentredString(center_x, y, line)
                y -= title_size * 1.2

        c.setFont("Times-Roman", text_size)
        y = page_height - 5.25 * inch
        for line in wrap_text(entry["Certificate_Text"], "Times-Roman", text_size, avail_width):
            c.drawCentredString(center_x, y, line)
            y -= text_size * 1.2

        c.setFont("Times-Roman", date_size)
        y = page_height - 8.25 * inch
        for line in entry["Formatted_Date"].split("\n"):
            c.drawCentredString(center_x, y, line)
            y -= date_size * 1.2

        right_x = page_width - right_margin
        c.setFont("Times-Roman", 12)
        y = page_height - 9.5 * inch
        c.drawRightString(right_x, y, "_____________________________________")
        c.setFont("Times-Roman", 14)
        c.drawRightString(right_x, y - 14 * 1.2, "Stan Ellis")
        c.drawRightString(right_x, y - 14 * 2.4, "Assemblyman, 32nd District")

    c.save()
    buffer.seek(0)
    return buffer.read()

approved_entries = [c for c in final_cert_rows if c.get("approved")]
if not approved_entries:
    st.error("No certificates were approved.")
else:
    doc = generate_word_certificates(approved_entries)
    tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_docx.name)
    tmp_docx.seek(0)
    if st.download_button(
        label="CreateCert",
        data=tmp_docx.read(),
        file_name="Certificates.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ):
        log_certificates(
            parsed_entries,
            approved_entries,
            pdf_text,
            source=source_type,
            global_comment=global_comment,
        )

    pdf_bytes = generate_pdf_certificates(approved_entries)
    if st.download_button(
        label="CreateCert PDF",
        data=pdf_bytes,
        file_name="Certificates.pdf",
        mime="application/pdf",
    ):
        log_certificates(
            parsed_entries,
            approved_entries,
            pdf_text,
            source=source_type,
            global_comment=global_comment,
        )
